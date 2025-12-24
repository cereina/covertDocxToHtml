#!/usr/bin/env python3
"""
DOCX → HTML converter with Word-style → WET-BOEW/GCWeb class mapping.

Goals (pragmatic):
- Preserve paragraph order, headings, basic inline emphasis (bold/italic),
  simple lists, and tables.
- Add GCWeb/WET classes based on Word paragraph styles (especially custom styles
  named like "WET Alert Success", "WET Lead", etc.).
- Emit clean HTML that can be dropped into a GCWeb page template.

Limitations:
- Word list numbering detection varies by authoring tool; this script uses a
  best-effort check for docx numbering properties and the built-in "List Bullet"
  / "List Number" styles.
- Complex layout (text boxes, multi-column sections, floating images) is not handled.

Usage:
  python docx_to_gcweb_html.py input.docx -o output.html
"""

from __future__ import annotations
import argparse
import html
from pathlib import Path
from typing import Optional, List, Tuple

from docx import Document
from docx.oxml.ns import qn

# ---------------- Style → HTML mapping ----------------
# Map Word paragraph style names to (tag, classes, wrapper_tag, wrapper_classes)
# wrapper_tag/classes let you wrap a paragraph inside a component container.
STYLE_MAP = {
    # Typography
    "WET Lead":           ("p",    "lead",              None, None),
    "WET Small":          ("p",    "small",             None, None),  # or render <small> inside
    "WET Muted":          ("p",    "text-muted",        None, None),
    "WET Blockquote":     ("blockquote", None,          None, None),

    # Alerts (GCWeb uses Bootstrap alerts)
    "WET Alert Success":  ("p",    None,               "section", "alert alert-success"),
    "WET Alert Info":     ("p",    None,               "section", "alert alert-info"),
    "WET Alert Warning":  ("p",    None,               "section", "alert alert-warning"),
    "WET Alert Danger":   ("p",    None,               "section", "alert alert-danger"),

    # Well
    "WET Well":           ("p",    None,               "div",     "well"),

    # Buttons (render as <a> placeholders; you can post-process into real links/buttons)
    "WET Button Primary": ("a",    "btn btn-primary",  None, None),
    "WET Button Default": ("a",    "btn btn-default",  None, None),
    "WET Button Danger":  ("a",    "btn btn-danger",   None, None),

    # Table marker paragraphs: the next table will inherit these classes
    "WET Table Basic":    ("p",    None,               None, None),
    "WET Table Striped":  ("p",    None,               None, None),
    "WET Table Bordered": ("p",    None,               None, None),
    "WET Table Hover":    ("p",    None,               None, None),
    "WET Table Condensed":("p",    None,               None, None),
}

TABLE_STYLE_TO_CLASS = {
    "WET Table Basic":     "table",
    "WET Table Striped":   "table table-striped",
    "WET Table Bordered":  "table table-bordered",
    "WET Table Hover":     "table table-hover",
    "WET Table Condensed": "table table-condensed",
}

# ---------------- Helpers ----------------
def esc(s: str) -> str:
    return html.escape(s, quote=True)

def paragraph_is_list(paragraph) -> Tuple[bool, Optional[str]]:
    """
    Returns (is_list, list_kind) where list_kind is "ul" or "ol".
    Best-effort detection: checks built-in style names and numbering properties.
    """
    style_name = getattr(paragraph.style, "name", "") or ""
    if style_name.startswith("List Bullet"):
        return True, "ul"
    if style_name.startswith("List Number"):
        return True, "ol"

    # Check numbering properties in XML
    p = paragraph._p
    pPr = p.pPr
    if pPr is not None and pPr.numPr is not None:
        # We can't reliably infer bullet vs number without inspecting numbering definitions,
        # but most numbered lists authored in Word use numPr too.
        # Default to ul unless style suggests numbering.
        return True, "ul"
    return False, None

def runs_to_html(paragraph) -> str:
    """
    Convert runs to basic inline HTML (bold/italic/underline).
    """
    parts: List[str] = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        frag = esc(text).replace("\n", "<br/>")
        if run.bold:
            frag = f"<strong>{frag}</strong>"
        if run.italic:
            frag = f"<em>{frag}</em>"
        if run.underline:
            frag = f"<u>{frag}</u>"
        parts.append(frag)
    # If there were no runs (rare), fall back to paragraph.text
    if not parts:
        return esc(paragraph.text)
    return "".join(parts)

def heading_level(style_name: str) -> Optional[int]:
    if style_name.startswith("Heading "):
        try:
            lvl = int(style_name.split("Heading ", 1)[1])
            if 1 <= lvl <= 6:
                return lvl
        except Exception:
            return None
    return None

def table_to_html(table, table_classes: str) -> str:
    rows = table.rows
    if not rows:
        return ""
    html_rows: List[str] = []
    # Assume first row is header if it looks like a header (simple heuristic: bold text in first row)
    header_cells = rows[0].cells
    ths = "".join(f"<th scope='col'>{esc(c.text.strip())}</th>" for c in header_cells)
    html_rows.append(f"<thead><tr>{ths}</tr></thead>")
    body_parts: List[str] = []
    for r in rows[1:]:
        tds = "".join(f"<td>{esc(c.text.strip())}</td>" for c in r.cells)
        body_parts.append(f"<tr>{tds}</tr>")
    html_rows.append("<tbody>" + "".join(body_parts) + "</tbody>")
    return f"<table class='{esc(table_classes)}'>" + "".join(html_rows) + "</table>"

# ---------------- Main conversion ----------------
def convert(docx_path: Path) -> str:
    doc = Document(str(docx_path))

    out: List[str] = []
    out.append("<main property='mainContentOfPage' class='container'>")

    current_list_kind: Optional[str] = None
    pending_table_class: Optional[str] = None

    # Iterate block items in order: paragraphs + tables
    # python-docx doesn't provide a direct unified iterator, so we walk the XML body children.
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag

        # Paragraph
        if tag.endswith("}p"):
            paragraph = None
            # find corresponding paragraph object
            for p in doc.paragraphs:
                if p._p is child:
                    paragraph = p
                    break
            if paragraph is None:
                continue

            style_name = getattr(paragraph.style, "name", "") or ""

            # Track table classes via marker paragraphs (immediately before a table)
            if style_name in TABLE_STYLE_TO_CLASS:
                pending_table_class = TABLE_STYLE_TO_CLASS[style_name]

            # Close list if we hit a non-list paragraph
            is_list, kind = paragraph_is_list(paragraph)
            if not is_list and current_list_kind:
                out.append(f"</{current_list_kind}>")
                current_list_kind = None

            # Headings
            lvl = heading_level(style_name)
            if lvl:
                out.append(f"<h{lvl}>{runs_to_html(paragraph)}</h{lvl}>")
                continue

            # List items
            if is_list:
                if current_list_kind != kind:
                    if current_list_kind:
                        out.append(f"</{current_list_kind}>")
                    current_list_kind = kind
                    out.append(f"<{current_list_kind}>")
                out.append(f"<li>{runs_to_html(paragraph)}</li>")
                continue

            # Style-mapped paragraph/component
            tag_name, classes, wrapper_tag, wrapper_classes = STYLE_MAP.get(
                style_name, ("p", None, None, None)
            )

            # Buttons as <a href="#">
            attrs = []
            if tag_name == "a":
                attrs.append("href='#'")
                attrs.append("role='button'")

            if classes:
                attrs.append(f"class='{esc(classes)}'")

            inner = runs_to_html(paragraph)
            element_html = f"<{tag_name} " + " ".join(attrs) + f">{inner}</{tag_name}>" if attrs else f"<{tag_name}>{inner}</{tag_name}>"

            if wrapper_tag:
                wc = f" class='{esc(wrapper_classes)}'" if wrapper_classes else ""
                element_html = f"<{wrapper_tag}{wc}>{element_html}</{wrapper_tag}>"

            out.append(element_html)

        # Table
        elif tag.endswith("}tbl"):
            table = None
            for t in doc.tables:
                if t._tbl is child:
                    table = t
                    break
            if table is None:
                continue
            # close any open list before tables
            if current_list_kind:
                out.append(f"</{current_list_kind}>")
                current_list_kind = None

            table_classes = pending_table_class or "table"
            pending_table_class = None
            out.append(table_to_html(table, table_classes))

        # Section properties etc are ignored

    if current_list_kind:
        out.append(f"</{current_list_kind}>")

    out.append("</main>")
    return "\n".join(out)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="Path to input .docx")
    ap.add_argument("-o", "--output", help="Output HTML file path. If omitted, prints to stdout.")
    args = ap.parse_args()

    html_out = convert(Path(args.input))

    if args.output:
        Path(args.output).write_text(html_out, encoding="utf-8")
    else:
        print(html_out)

if __name__ == "__main__":
    main()
