#!/usr/bin/env python3
"""
DOCX → HTML converter with Word-style → WET-BOEW/GCWeb class mapping.

Now supports:
- <details>/<summary> blocks via Word styles:
    * "WET Details Summary"
    * "WET Details Content"
- GCWeb/WET accordion via Word styles:
    * "WET Accordion Start"
    * "WET Accordion Heading"
    * "WET Accordion Panel"
    * "WET Accordion End"
  Output structure:
    <section class="wb-accordion">
      <details><summary>...</summary><p>...</p></details>
      ...
    </section>
- Pagination via Word styles:
    * "WET Pagination Start"
    * "WET Pagination Item"
    * "WET Pagination Active"
    * "WET Pagination Disabled"
    * "WET Pagination End"
  Output structure:
    <nav aria-label="Pagination">
      <ul class="pagination"> ... </ul>
    </nav>

Usage:
  python docx_to_gcweb_html.py input.docx -o output.html
"""

from __future__ import annotations
import argparse
import html
from pathlib import Path
from typing import Optional, List, Tuple

from docx import Document

# ---------------- Style → HTML mapping ----------------
# Map Word paragraph style names to (tag, classes, wrapper_tag, wrapper_classes)
STYLE_MAP = {
    # Typography
    "WET Lead":           ("p",    "lead",              None, None),
    "WET Small":          ("p",    "small",             None, None),
    "WET Muted":          ("p",    "text-muted",        None, None),
    "WET Blockquote":     ("blockquote", None,          None, None),

    # Alerts (GCWeb uses Bootstrap alerts)
    "WET Alert Success":  ("p",    None,               "section", "alert alert-success"),
    "WET Alert Info":     ("p",    None,               "section", "alert alert-info"),
    "WET Alert Warning":  ("p",    None,               "section", "alert alert-warning"),
    "WET Alert Danger":   ("p",    None,               "section", "alert alert-danger"),

    # Well
    "WET Well":           ("p",    None,               "div",     "well"),

    # Buttons (render as <a> placeholders; you can post-process into real links/buttons)
    "WET Button Primary": ("a",    "btn btn-primary",  None, None),
    "WET Button Default": ("a",    "btn btn-default",  None, None),
    "WET Button Danger":  ("a",    "btn btn-danger",   None, None),
    "WET Button Link":    ("a",    "btn btn-link",     None, None),

    # Markers handled by state machines (details/accordion/pagination) are NOT listed here.
    # Table marker paragraphs are also handled separately.
}

TABLE_STYLE_TO_CLASS = {
    "WET Table Basic":     "table",
    "WET Table Striped":   "table table-striped",
    "WET Table Bordered":  "table table-bordered",
    "WET Table Hover":     "table table-hover",
    "WET Table Condensed": "table table-condensed",
}

# ---------------- Helpers ----------------
def esc(s: str) -> str:
    return html.escape(s, quote=True)

def paragraph_is_list(paragraph) -> Tuple[bool, Optional[str]]:
    """
    Returns (is_list, list_kind) where list_kind is "ul" or "ol".
    Best-effort detection: checks built-in style names and numbering properties.
    """
    style_name = getattr(paragraph.style, "name", "") or ""
    if style_name.startswith("List Bullet"):
        return True, "ul"
    if style_name.startswith("List Number"):
        return True, "ol"

    # Check numbering properties in XML (best-effort; type unknown => default ul)
    pPr = paragraph._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True, "ul"
    return False, None

def runs_to_html(paragraph) -> str:
    """Convert runs to basic inline HTML (bold/italic/underline)."""
    parts: List[str] = []
    for run in paragraph.runs:
        if not run.text:
            continue
        frag = esc(run.text).replace("\n", "<br/>")
        if run.bold:
            frag = f"<strong>{frag}</strong>"
        if run.italic:
            frag = f"<em>{frag}</em>"
        if run.underline:
            frag = f"<u>{frag}</u>"
        parts.append(frag)
    return "".join(parts) if parts else esc(paragraph.text)

def heading_level(style_name: str) -> Optional[int]:
    if style_name.startswith("Heading "):
        try:
            lvl = int(style_name.split("Heading ", 1)[1])
            if 1 <= lvl <= 6:
                return lvl
        except Exception:
            return None
    return None

def table_to_html(table, table_classes: str) -> str:
    rows = table.rows
    if not rows:
        return ""
    html_rows: List[str] = []

    header_cells = rows[0].cells
    ths = "".join(f"<th scope='col'>{esc(c.text.strip())}</th>" for c in header_cells)
    html_rows.append(f"<thead><tr>{ths}</tr></thead>")

    body_parts: List[str] = []
    for r in rows[1:]:
        tds = "".join(f"<td>{esc(c.text.strip())}</td>" for c in r.cells)
        body_parts.append(f"<tr>{tds}</tr>")
    html_rows.append("<tbody>" + "".join(body_parts) + "</tbody>")

    return f"<table class='{esc(table_classes)}'>" + "".join(html_rows) + "</table>"

# ---------------- Main conversion ----------------
def convert(docx_path: Path) -> str:
    doc = Document(str(docx_path))

    out: List[str] = []
    out.append("<main property='mainContentOfPage' class='container'>")

    # Fast lookup: underlying XML element -> paragraph/table object
    para_by_el = {p._p: p for p in doc.paragraphs}
    tbl_by_el = {t._tbl: t for t in doc.tables}

    # Generic list state
    current_list_kind: Optional[str] = None

    # Table marker state
    pending_table_class: Optional[str] = None
    pending_table_responsive: bool = False  # if True wrap next table in <div class="table-responsive">

    # Details state (outside accordion)
    in_details: bool = False

    # Accordion state
    in_accordion: bool = False
    accordion_item_open: bool = False  # <details> inside accordion open?

    # Pagination state
    in_pagination: bool = False

    def close_list_if_open():
        nonlocal current_list_kind
        if current_list_kind:
            out.append(f"</{current_list_kind}>")
            current_list_kind = None

    def close_details_if_open():
        nonlocal in_details
        if in_details:
            out.append("</details>")
            in_details = False

    def close_accordion_item_if_open():
        nonlocal accordion_item_open
        if accordion_item_open:
            out.append("</details>")
            accordion_item_open = False

    def close_accordion_if_open():
        nonlocal in_accordion
        if in_accordion:
            close_accordion_item_if_open()
            out.append("</section>")
            in_accordion = False

    def close_pagination_if_open():
        nonlocal in_pagination
        if in_pagination:
            out.append("</ul></nav>")
            in_pagination = False

    def close_all_structures_before_new_block():
        # When entering a heading/table/normal paragraph, close incompatible open blocks.
        # Note: Accordion is explicit start/end; we don't auto-close it unless we hit its End marker.
        close_list_if_open()
        # Details blocks auto-close when content stops, but we also close when "something else" starts.
        close_details_if_open()
        # Pagination is explicit start/end; don't auto-close unless we hit End marker, but tables/headings should close it.
        # We'll be conservative: close pagination when a table appears (common in authored docs).
        # (You can remove this if you want strict marker-only behaviour.)
        # close_pagination_if_open()

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag

        # Paragraph
        if tag.endswith("}p"):
            paragraph = para_by_el.get(child)
            if paragraph is None:
                continue

            style_name = getattr(paragraph.style, "name", "") or ""
            text_html = runs_to_html(paragraph)

            # ---------- Accordion markers ----------
            if style_name == "WET Accordion Start":
                close_list_if_open()
                close_details_if_open()
                close_pagination_if_open()
                close_accordion_if_open()
                out.append("<section class='wb-accordion'>")
                in_accordion = True
                accordion_item_open = False
                continue

            if style_name == "WET Accordion End":
                close_list_if_open()
                close_details_if_open()
                close_pagination_if_open()
                close_accordion_if_open()
                continue

            if in_accordion:
                if style_name == "WET Accordion Heading":
                    close_list_if_open()
                    close_details_if_open()
                    # start a new accordion item
                    close_accordion_item_if_open()
                    out.append("<details>")
                    out.append(f"<summary>{text_html}</summary>")
                    accordion_item_open = True
                    continue
                if style_name == "WET Accordion Panel":
                    # panels become paragraphs inside the current accordion item
                    if not accordion_item_open:
                        # If author forgot a heading, create a fallback item
                        out.append("<details><summary>Details</summary>")
                        accordion_item_open = True
                    out.append(f"<p>{text_html}</p>")
                    continue
                # If any other paragraph appears inside accordion, treat it as panel content (safe default)
                if accordion_item_open:
                    out.append(f"<p>{text_html}</p>")
                    continue

            # ---------- Pagination markers ----------
            if style_name == "WET Pagination Start":
                close_list_if_open()
                close_details_if_open()
                close_pagination_if_open()
                # Keep accordion separate; if someone starts pagination inside accordion, it's invalid—treat as outside.
                out.append("<nav aria-label='Pagination'><ul class='pagination'>")
                in_pagination = True
                continue

            if style_name == "WET Pagination End":
                close_list_if_open()
                close_details_if_open()
                close_pagination_if_open()
                continue

            if in_pagination:
                close_list_if_open()
                close_details_if_open()

                label = text_html
                if style_name == "WET Pagination Active":
                    out.append(f"<li class='active'><a href='#' aria-current='page'>{label}</a></li>")
                    continue
                if style_name == "WET Pagination Disabled":
                    out.append(f"<li class='disabled'><span>{label}</span></li>")
                    continue
                # default item
                if style_name in ("WET Pagination Item",):
                    out.append(f"<li><a href='#'>{label}</a></li>")
                    continue

                # If a non-pagination paragraph appears, close pagination and fall through
                close_pagination_if_open()
                # fall through to normal handling of this paragraph

            # ---------- Details/Summary blocks (outside accordion) ----------
            if style_name == "WET Details Summary":
                close_list_if_open()
                close_details_if_open()
                out.append("<details>")
                out.append(f"<summary>{text_html}</summary>")
                in_details = True
                continue

            if style_name == "WET Details Content":
                close_list_if_open()
                if not in_details:
                    # If author forgot summary, open a default details
                    out.append("<details><summary>Details</summary>")
                    in_details = True
                out.append(f"<p>{text_html}</p>")
                continue

            # If details content has ended, close it before continuing
            if in_details:
                close_details_if_open()

            # ---------- Table marker paragraphs ----------
            if style_name in TABLE_STYLE_TO_CLASS:
                pending_table_class = TABLE_STYLE_TO_CLASS[style_name]
            if style_name == "WET Table Responsive":
                pending_table_responsive = True

            # ---------- Close lists when necessary ----------
            is_list, kind = paragraph_is_list(paragraph)
            if not is_list and current_list_kind:
                out.append(f"</{current_list_kind}>")
                current_list_kind = None

            # ---------- Headings ----------
            lvl = heading_level(style_name)
            if lvl:
                close_list_if_open()
                close_details_if_open()
                # headings should not live inside pagination
                # (leave accordion as explicit markers)
                out.append(f"<h{lvl}>{text_html}</h{lvl}>")
                continue

            # ---------- Lists ----------
            if is_list:
                # Enter list
                if current_list_kind != kind:
                    if current_list_kind:
                        out.append(f"</{current_list_kind}>")
                    current_list_kind = kind
                    out.append(f"<{current_list_kind}>")
                out.append(f"<li>{text_html}</li>")
                continue

            # ---------- Regular mapped paragraph/component ----------
            tag_name, classes, wrapper_tag, wrapper_classes = STYLE_MAP.get(
                style_name, ("p", None, None, None)
            )

            attrs: List[str] = []
            if tag_name == "a":
                attrs.append("href='#'")
                attrs.append("role='button'")
            if classes:
                attrs.append(f"class='{esc(classes)}'")

            if attrs:
                element_html = f"<{tag_name} " + " ".join(attrs) + f">{text_html}</{tag_name}>"
            else:
                element_html = f"<{tag_name}>{text_html}</{tag_name}>"

            if wrapper_tag:
                wc = f" class='{esc(wrapper_classes)}'" if wrapper_classes else ""
                element_html = f"<{wrapper_tag}{wc}>{element_html}</{wrapper_tag}>"

            out.append(element_html)

        # Table
        elif tag.endswith("}tbl"):
            table = tbl_by_el.get(child)
            if table is None:
                continue

            close_list_if_open()
            close_details_if_open()
            # tables often shouldn't appear inside pagination
            close_pagination_if_open()
            # accordion remains explicit via end marker (don't auto-close)

            table_classes = pending_table_class or "table"
            pending_table_class = None

            table_html = table_to_html(table, table_classes)

            if pending_table_responsive:
                out.append(f"<div class='table-responsive'>{table_html}</div>")
                pending_table_responsive = False
            else:
                out.append(table_html)

        # Ignore other child types

    # Close any open structures
    close_list_if_open()
    close_details_if_open()
    close_pagination_if_open()
    close_accordion_if_open()

    out.append("</main>")
    return "\n".join(out)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="Path to input .docx")
    ap.add_argument("-o", "--output", help="Output HTML file path. If omitted, prints to stdout.")
    args = ap.parse_args()

    html_out = convert(Path(args.input))

    if args.output:
        Path(args.output).write_text(html_out, encoding="utf-8")
    else:
        print(html_out)

if __name__ == "__main__":
    main()
